from __future__ import annotations
import logging
import os
import re
import shutil
import subprocess
import tempfile
from typing import Any, Callable
import boto3
import cv2
import easyocr
import fitz
import numpy as np
import requests
from botocore.exceptions import BotoCoreError, ClientError
from docx import Document
from app.core.config import (
    AWS_S3_ACCESS_KEY,
    AWS_S3_REGION,
    AWS_S3_SECRET_KEY,
    LAPLACIAN_VAR_THRESHOLD,
    MIN_CONTRAST_THRESHOLD,
    MIN_IMAGE_HEIGHT,
    MIN_IMAGE_WIDTH,
    OCR_USE_GPU,
    TEMP_DIR,
    VALIDATE_ALL_PDF_PAGES,
)
from app.models.validation import ValidationCheck, ValidationReport

logger = logging.getLogger(__name__)

S3_URI = re.compile(r"^s3://([^/]+)/(.+)$")

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
DOC_MIME = "application/msword"
TXT_MIME = "text/plain"
PDF_MIME = "application/pdf"
IMAGE_EXTENSIONS = (".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".webp")

TEXT_FILE_ENCODINGS = ("utf-8", "utf-8-sig", "cp1258", "latin-1")
SUPPORTED_OCR_TYPES = (
    "image/*",
    PDF_MIME,
    TXT_MIME,
    DOC_MIME,
    DOCX_MIME,
)

MIME_TO_EXT: dict[str, str] = {
    "image/png": "png",
    "image/jpeg": "jpg",
    "image/jpg": "jpg",
    PDF_MIME: "pdf",
    DOCX_MIME: "docx",
    DOC_MIME: "doc",
    TXT_MIME: "txt",
}


class ImageTooBlurryError(Exception):
    """Raised when Laplacian variance is below threshold."""


class UnsupportedFileTypeError(Exception):
    """Raised when file type is not supported by OCR pipeline."""


def safe_unlink(path: str) -> None:
    try:
        if path and os.path.isfile(path):
            os.unlink(path)
    except OSError:
        pass


def suffix_for_type(file_type: str) -> str:
    ft = (file_type or "").lower()
    for mime, ext in MIME_TO_EXT.items():
        if mime in ft:
            return ext
    return "bin"


def resolve_content_type(file_path: str, file_type: str) -> str:
    ft = (file_type or "").lower()
    lp = file_path.lower()
    if TXT_MIME in ft or lp.endswith(".txt"):
        return "txt"
    if DOCX_MIME in ft or lp.endswith(".docx"):
        return "docx"
    if DOC_MIME in ft or lp.endswith(".doc"):
        return "doc"
    if "pdf" in ft or lp.endswith(".pdf"):
        return "pdf"
    if ft.startswith("image/") or lp.endswith(IMAGE_EXTENSIONS):
        return "image"
    return "unsupported"


class FileDownloader:
    """Downloads a remote file (S3 or HTTP) to a local temp path."""

    def __init__(self, temp_dir: str = TEMP_DIR) -> None:
        self.temp_dir = temp_dir

    def download(self, file_url: str, file_type: str) -> str:
        ext = suffix_for_type(file_type)
        fd, path = tempfile.mkstemp(suffix=f".{ext}", dir=self.temp_dir)
        os.close(fd)
        try:
            if file_url.startswith("s3://"):
                self.from_s3(file_url, path)
            else:
                self.from_http(file_url, path)
        except Exception:
            safe_unlink(path)
            raise
        return path

    def from_s3(self, file_url: str, path: str) -> None:
        m = S3_URI.match(file_url)
        if not m:
            raise ValueError(f"Invalid S3 URI: {file_url!r}")
        bucket, key = m.group(1), m.group(2)
        try:
            boto3.client(
                "s3",
                aws_access_key_id=AWS_S3_ACCESS_KEY or None,
                aws_secret_access_key=AWS_S3_SECRET_KEY or None,
                region_name=AWS_S3_REGION,
            ).download_file(bucket, key, path)
        except (ClientError, BotoCoreError):
            logger.exception("S3 download failed bucket=%s key=%s", bucket, key)
            raise

    @staticmethod
    def from_http(file_url: str, path: str) -> None:
        resp = requests.get(file_url, timeout=60, stream=True)
        resp.raise_for_status()
        with open(path, "wb") as f:
            for chunk in resp.iter_content(chunk_size=8192):
                if chunk:
                    f.write(chunk)


class ImageAnalyzer:
    """Low-level image operations: blur check, contrast, OCR via EasyOCR."""

    PDF_RASTER_MATRIX = fitz.Matrix(150 / 72, 150 / 72)

    def __init__(
        self,
        languages: list[str] | None = None,
        blur_threshold: float = LAPLACIAN_VAR_THRESHOLD,
        temp_dir: str = TEMP_DIR,
        use_gpu: bool = OCR_USE_GPU,
    ) -> None:
        self.languages = languages or ["vi", "en"]
        self.blur_threshold = blur_threshold
        self.temp_dir = temp_dir
        self.use_gpu = use_gpu

    @property
    def reader(self) -> easyocr.Reader:
        if self.reader is None:
            logger.info("Initializing EasyOCR (gpu=%s)", self.use_gpu)
            self.reader = easyocr.Reader(self.languages, gpu=self.use_gpu)
        return self.reader

    def ocr_file(self, image_path: str) -> str:
        results = self.reader.readtext(image_path)
        return self.join(results)

    def ocr_array(self, img_bgr: np.ndarray) -> str:
        img_rgb = cv2.cvtColor(img_bgr, cv2.COLOR_BGR2RGB)
        results = self.reader.readtext(img_rgb)
        return self.join(results)

    @staticmethod
    def join(results: list[Any]) -> str:
        return " ".join(str(r[1]) for r in results if len(r) > 1).strip()

    def laplacian_variance(self, img_bgr: np.ndarray) -> float:
        gray = (
            cv2.cvtColor(img_bgr, cv2.COLOR_BGR2GRAY)
            if len(img_bgr.shape) == 3
            else img_bgr
        )
        return float(cv2.Laplacian(gray, cv2.CV_64F).var())

    def ensure_not_blurry(self, image_path: str, context: str = "Image") -> None:
        img = cv2.imread(image_path)
        if img is None:
            raise ImageTooBlurryError(f"{context}: cannot read image file")
        var = self.laplacian_variance(img)
        if var < self.blur_threshold:
            raise ImageTooBlurryError(
                f"{context} too blurry (variance={var:.1f} < {self.blur_threshold})"
            )

    def pdf_to_bgr_pages(self, pdf_path: str) -> list[np.ndarray]:
        with fitz.open(pdf_path) as doc:
            return [self.page_to_bgr(doc, i) for i in range(len(doc))]

    @classmethod
    def page_to_bgr(cls, doc: fitz.Document, index: int) -> np.ndarray:
        pix = doc.load_page(index).get_pixmap(matrix=cls.PDF_RASTER_MATRIX, alpha=False)
        rgb = fitz.Pixmap(fitz.csRGB, pix)
        img = np.frombuffer(rgb.samples, dtype=np.uint8).reshape(
            rgb.height, rgb.width, 3
        )
        return cv2.cvtColor(img, cv2.COLOR_RGB2BGR)

    def image_quality_checks(
        self, img_bgr: np.ndarray, label: str = "Image"
    ) -> list[ValidationCheck]:
        checks: list[ValidationCheck] = []
        h, w = img_bgr.shape[:2]

        if w < MIN_IMAGE_WIDTH or h < MIN_IMAGE_HEIGHT:
            checks.append(
                ValidationCheck(
                    name="resolution",
                    passed=False,
                    message=f"{label} resolution too low ({w}x{h}), min {MIN_IMAGE_WIDTH}x{MIN_IMAGE_HEIGHT}",
                )
            )

        gray = (
            cv2.cvtColor(img_bgr, cv2.COLOR_BGR2GRAY)
            if len(img_bgr.shape) == 3
            else img_bgr
        )
        blur_var = self.laplacian_variance(img_bgr)
        if blur_var < self.blur_threshold:
            checks.append(
                ValidationCheck(
                    name="blur",
                    passed=False,
                    message=f"{label} too blurry (variance={blur_var:.1f}, min={self.blur_threshold})",
                )
            )

        contrast = float(np.std(gray))
        if contrast < MIN_CONTRAST_THRESHOLD:
            checks.append(
                ValidationCheck(
                    name="contrast",
                    passed=False,
                    message=f"{label} contrast too low (std={contrast:.1f}, min={MIN_CONTRAST_THRESHOLD})",
                )
            )

        return checks


class FileValidator:
    """Runs quality checks and returns a ValidationReport."""

    def __init__(self, analyzer: ImageAnalyzer) -> None:
        self.analyzer = analyzer

    def validate(self, file_path: str, file_type: str, doc_id: int) -> ValidationReport:
        content_type = resolve_content_type(file_path, file_type)
        checks = self.run_checks(file_path, file_type, content_type)
        overall = not checks or all(c.passed for c in checks)

        if not overall:
            failed = [c.message or c.name for c in checks if not c.passed]
            logger.warning("Validation FAILED doc_id=%s: %s", doc_id, "; ".join(failed))

        return ValidationReport(
            doc_id=doc_id,
            file_type=file_type or "unknown",
            checks=checks,
            overall_passed=overall,
        )

    def run_checks(
        self, file_path: str, file_type: str, content_type: str
    ) -> list[ValidationCheck]:
        if content_type == "unsupported":
            return [
                ValidationCheck(
                    name="format", passed=False, message=f"Unsupported: {file_type}"
                )
            ]

        handlers: dict[str, Callable[[str], list[ValidationCheck]]] = {
            "image": self.check_image,
            "pdf": self.check_pdf,
            "docx": self.check_docx,
            "txt": self.check_text,
            "doc": self.check_text,
        }
        return handlers.get(content_type, lambda _: [])(file_path)

    def check_image(self, path: str) -> list[ValidationCheck]:
        img = cv2.imread(path)
        if img is None:
            return [
                ValidationCheck(
                    name="integrity", passed=False, message="Cannot read image file"
                )
            ]
        return self.analyzer.image_quality_checks(img)

    def check_pdf(self, path: str) -> list[ValidationCheck]:
        try:
            doc = fitz.open(path)
        except Exception as exc:
            return [
                ValidationCheck(
                    name="integrity", passed=False, message=f"PDF corrupt: {exc}"
                )
            ]

        if doc.page_count == 0:
            doc.close()
            return [
                ValidationCheck(
                    name="integrity", passed=False, message="PDF has no pages"
                )
            ]

        pages = range(len(doc)) if VALIDATE_ALL_PDF_PAGES else range(min(1, len(doc)))
        checks = []
        for i in pages:
            checks.extend(
                self.analyzer.image_quality_checks(
                    self.analyzer.page_to_bgr(doc, i), label=f"Page {i + 1}"
                )
            )
        doc.close()
        return checks

    @staticmethod
    def check_docx(path: str) -> list[ValidationCheck]:
        try:
            Document(path)
            return []
        except Exception as exc:
            return [
                ValidationCheck(
                    name="integrity", passed=False, message=f"DOCX corrupt: {exc}"
                )
            ]

    @staticmethod
    def check_text(path: str) -> list[ValidationCheck]:
        if os.path.getsize(path) == 0:
            return [
                ValidationCheck(name="integrity", passed=False, message="File is empty")
            ]
        return []


class TextExtractor:
    """Extracts plain text from supported file types."""

    def __init__(self, analyzer: ImageAnalyzer, temp_dir: str = TEMP_DIR) -> None:
        self.analyzer = analyzer
        self.temp_dir = temp_dir

    def extract(self, file_path: str, file_type: str) -> str:
        content_type = resolve_content_type(file_path, file_type)
        if content_type == "unsupported":
            raise UnsupportedFileTypeError(
                f"Unsupported file type: {file_type or 'application/octet-stream'}. "
                f"Supported: {', '.join(SUPPORTED_OCR_TYPES)}."
            )
        handlers: dict[str, Callable[[str], str]] = {
            "txt": self.read_text_file,
            "docx": self.extract_docx_text,
            "doc": self.extract_doc,
            "pdf": self.extract_pdf,
            "image": self.extract_image,
        }
        return handlers[content_type](file_path)

    def read_text_file(self, file_path: str) -> str:
        for enc in TEXT_FILE_ENCODINGS:
            try:
                with open(file_path, encoding=enc) as f:
                    return f.read().strip()
            except UnicodeDecodeError:
                continue
        with open(file_path, "rb") as f:
            return f.read().decode("latin-1", errors="ignore").strip()

    def extract_docx_text(self, file_path: str) -> str:
        doc = Document(file_path)
        return "\n".join(
            p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()
        ).strip()

    def extract_image(self, file_path: str) -> str:
        self.analyzer.ensure_not_blurry(file_path)
        return self.analyzer.ocr_file(file_path)

    def extract_pdf(self, file_path: str) -> str:
        pages = self.analyzer.pdf_to_bgr_pages(file_path)
        if not pages:
            return ""

        fd, first_path = tempfile.mkstemp(suffix=".png", dir=self.temp_dir)
        try:
            os.close(fd)
            cv2.imwrite(first_path, pages[0])
            self.analyzer.ensure_not_blurry(first_path, context="First PDF page")
        finally:
            safe_unlink(first_path)

        return "\n".join(self.analyzer.ocr_array(p) for p in pages).strip()

    def extract_doc(self, file_path: str) -> str:
        soffice = shutil.which("soffice")
        if not soffice:
            raise RuntimeError("LibreOffice (soffice) is not installed.")

        out_dir = tempfile.mkdtemp(dir=self.temp_dir)
        base = os.path.splitext(os.path.basename(file_path))[0]
        txt_path = os.path.join(out_dir, f"{base}.txt")
        try:
            subprocess.run(
                [
                    soffice,
                    "--headless",
                    "--convert-to",
                    "txt:Text",
                    "--outdir",
                    out_dir,
                    file_path,
                ],
                check=True,
                capture_output=True,
                text=True,
                timeout=120,
            )
            if not os.path.isfile(txt_path):
                raise RuntimeError(
                    "LibreOffice conversion succeeded but output file missing."
                )
            return self.read_text_file(txt_path)
        except subprocess.TimeoutExpired as exc:
            raise RuntimeError("Timed out converting .doc with LibreOffice.") from exc
        except subprocess.CalledProcessError as exc:
            raise RuntimeError(
                f"LibreOffice conversion failed: {(exc.stderr or '').strip()}"
            ) from exc
        finally:
            safe_unlink(txt_path)
            try:
                os.rmdir(out_dir)
            except OSError:
                pass


class OcrService:
    def __init__(
        self,
        languages: list[str] | None = None,
        blur_threshold: float = LAPLACIAN_VAR_THRESHOLD,
        temp_dir: str = TEMP_DIR,
        use_gpu: bool = OCR_USE_GPU,
    ) -> None:
        analyzer = ImageAnalyzer(languages, blur_threshold, temp_dir, use_gpu)
        self.downloader = FileDownloader(temp_dir)
        self.extractor = TextExtractor(analyzer, temp_dir)
        self.validator = FileValidator(analyzer)

    def download_to_temp(self, file_url: str, file_type: str) -> str:
        return self.downloader.download(file_url, file_type)

    def validate(self, file_path: str, file_type: str, doc_id: int) -> ValidationReport:
        return self.validator.validate(file_path, file_type, doc_id)

    def run_ocr(self, file_path: str, file_type: str) -> str:
        return self.extractor.extract(file_path, file_type)

    @staticmethod
    def cleanup_temp(path: str) -> None:
        safe_unlink(path)


ocr_service = OcrService()
